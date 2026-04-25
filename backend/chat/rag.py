import csv
import hashlib
import json
import os
import re
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.documents import Document
from langchain_classic.chains import create_retrieval_chain
from langchain_classic.chains.combine_documents import create_stuff_documents_chain
from langchain_community.callbacks import get_openai_callback
from openai import APIConnectionError, APITimeoutError
from PyPDF2 import PdfReader
from typing import List, Dict, Tuple, Any
from pathlib import Path

NO_ANSWER_FALLBACK = "I don't have information about this in the provided documents."
UPSTREAM_ERROR_MSG = "Couldn't reach the AI provider. Check your network or VPN and try again."
GENERIC_ERROR_MSG = "Something went wrong generating a response. Please try again."


class UpstreamUnavailable(Exception):
    """Raised when the LLM/embedding provider is unreachable (network / proxy)."""

def _openai_kwargs() -> dict:
    """Return extra kwargs for OpenAI clients (e.g. proxy support via OPENAI_PROXY env var).

    Sets a connect timeout so a blocked api.openai.com fails fast (~5s) instead of
    holding each request for the full OS TCP retry window (~60s)."""
    import httpx
    timeout = httpx.Timeout(connect=5.0, read=60.0, write=10.0, pool=5.0)
    proxy = os.environ.get("OPENAI_PROXY", "").strip()
    if proxy:
        return {
            "http_client": httpx.Client(proxy=proxy, timeout=timeout),
            "http_async_client": httpx.AsyncClient(proxy=proxy, timeout=timeout),
        }
    return {
        "http_client": httpx.Client(timeout=timeout),
        "http_async_client": httpx.AsyncClient(timeout=timeout),
    }

CHAT_PROMPT_TEMPLATE = """
You are a helpful AI assistant with two sources of context:

1. PROVIDED DOCUMENTS — the authoritative source for factual answers about the subject matter.
2. PERSONAL CONTEXT — what you know about this specific user. Use it to personalize tone, pick
   relevant examples, and to answer direct questions the user asks about themselves.

Rules:
- Factual subject-matter claims MUST be grounded in the document context. Do not invent facts.
- If the documents cannot answer a factual question about the subject matter, respond with EXACTLY:
  I don't have information about this in the provided documents.
- You MAY answer questions the user asks about themselves (\"who am I\", \"what am I working on\")
  using the personal context below, even when the documents don't cover it.

About the user:
{user_profile}

What you remember about this user:
{user_memories}

Conversation history:
{chat_history}

Document context:
{context}

User question:
{input}

Answer clearly and concisely.
"""


CONVERSATIONAL_PROMPT_TEMPLATE = """
You are DocuChat, a helpful AI assistant for chatting with uploaded documents.
Reply briefly and warmly to greetings, capability questions, and small talk.
If asked what you can do, mention you can answer questions about the user's
uploaded documents and remember things they tell you about themselves.

Conversation history:
{chat_history}

User message:
{input}

Reply concisely.
"""


PERSONAL_PROMPT_TEMPLATE = """
The user is asking a question about themselves. Answer using ONLY the profile
and memories below. If neither contains the answer, say so plainly — for
example: "I don't know that about you yet — feel free to tell me."
Do NOT respond with the document-fallback sentence; this is not a document
question.

About the user:
{user_profile}

What you remember about this user:
{user_memories}

Conversation history:
{chat_history}

User question:
{input}

Answer briefly and personally.
"""


# Default to 'factual' on ambiguity so we never strip grounding from real questions.
_GREETING_START_RE = re.compile(
    r"^\s*(hi+|hello+|hey+|yo+|sup|good\s+(morning|afternoon|evening|day)|"
    r"howdy|greetings|thanks?|thank\s+you|thx|ty|bye|goodbye|cheers)\b",
    re.IGNORECASE,
)

_META_PATTERNS = [
    re.compile(r"\b(can|could|would|will)\s+you\s+(help|assist)\b", re.I),
    re.compile(r"\bwhat\s+can\s+you\s+(do|help\s+with)\b", re.I),
    re.compile(r"\bwho\s+are\s+you\b", re.I),
    re.compile(r"\bhow\s+(do|does)\s+(you|this)\s+work\b", re.I),
    re.compile(r"\bhow\s+are\s+you\b", re.I),
]

_PERSONAL_PATTERNS = [
    re.compile(r"\bwhat\s+do\s+you\s+(know|remember)\s+about\s+me\b", re.I),
    re.compile(r"\bwho\s+am\s+i\b", re.I),
    re.compile(r"\bwhat(?:'?s|\s+is)\s+my\s+name\b", re.I),
    re.compile(r"\bwhat\s+(am|was)\s+i\s+(working|building|doing)\b", re.I),
    re.compile(r"\bremember\s+about\s+me\b", re.I),
    re.compile(r"\babout\s+myself\b", re.I),
    re.compile(r"\bmy\s+(role|job|profession|background)\b", re.I),
]


def classify_intent(message: str) -> str:
    """Return one of 'conversational', 'personal', or 'factual'."""
    text = (message or "").strip()
    if not text:
        return "conversational"

    if any(p.search(text) for p in _PERSONAL_PATTERNS):
        return "personal"

    if _GREETING_START_RE.search(text) and len(text.split()) <= 8:
        return "conversational"

    if any(p.search(text) for p in _META_PATTERNS):
        return "conversational"

    return "factual"


def _format_memory_block(memories: List[str]) -> str:
    if not memories:
        return "(no memories yet)"
    return "\n".join(f"- {m}" for m in memories)


def _format_profile(profile_text: str) -> str:
    return profile_text.strip() if profile_text and profile_text.strip() else "(no profile set)"


class RAGSystem:
    """RAG (Retrieval-Augmented Generation) system for chat"""

    def __init__(self):
        kwargs = _openai_kwargs()
        # max_retries=0 + explicit timeout so a blocked upstream errors out in
        # ~5-10s instead of the SDK's default 3-attempt × OS-TCP-retry chain
        # (which can be 60-180s on Windows).
        self.embeddings = OpenAIEmbeddings(
            model="text-embedding-3-small",
            max_retries=0,
            timeout=10.0,
            **kwargs,
        )
        self.llm = ChatOpenAI(
            model="gpt-4o-mini",
            temperature=0.0,
            max_retries=0,
            timeout=30.0,
            **kwargs,
        )
        self.llm_streaming = ChatOpenAI(
            model="gpt-4o-mini",
            temperature=0.0,
            streaming=True,
            max_retries=0,
            timeout=30.0,
            **kwargs,
        )
        self.kb_path = Path(__file__).parent.parent / "knowledge_base"
        self.vector_store_dir = Path(__file__).parent.parent / "vector_store"
        self.vector_store_path = self.vector_store_dir / "faiss_index"
        self.manifest_path = self.vector_store_dir / "manifest.json"
        self.vector_store = None
        self.retriever = None
        self.failed_files: List[Tuple[str, str]] = []
        self._initialize_knowledge_base()

    # -------------------------
    # Knowledge base setup
    # -------------------------
    def _initialize_knowledge_base(self):
        self.kb_path.mkdir(exist_ok=True)
        self.vector_store_dir.mkdir(exist_ok=True)
        current_manifest = self._build_kb_manifest()

        if self._can_use_cached_index(current_manifest):
            self.vector_store = FAISS.load_local(
                str(self.vector_store_path),
                self.embeddings,
                allow_dangerous_deserialization=True,
            )
            print("Loaded cached vector store from disk.")
        else:
            print("Rebuilding vector store (knowledge base changed or cache missing).")
            documents = self._load_knowledge_documents()

            if not documents:
                dummy_doc = Document(
                    page_content="This is a demo RAG system.",
                    metadata={"source": "demo", "page": 1}
                )
                self.vector_store = FAISS.from_documents(
                    [dummy_doc], self.embeddings
                )
            else:
                splitter = RecursiveCharacterTextSplitter(
                    chunk_size=1000,
                    chunk_overlap=200
                )
                chunks = splitter.split_documents(documents)
                self.vector_store = FAISS.from_documents(
                    chunks, self.embeddings
                )

            self.vector_store.save_local(str(self.vector_store_path))
            self._write_manifest(current_manifest)
            print("Vector store saved to disk.")

        self.retriever = self.vector_store.as_retriever(
            search_kwargs={"k": 3}
        )

    def _load_knowledge_documents(self) -> List[Document]:
        documents = []
        self.failed_files: List[Tuple[str, str]] = []

        for file_path in self.kb_path.iterdir():
            if not file_path.is_file():
                continue

            ext = file_path.suffix.lower()
            supported = {".pdf", ".txt", ".md", ".docx", ".csv"}
            if ext not in supported:
                continue

            try:
                if ext == ".pdf":
                    loaded = self._load_pdf(file_path)
                elif ext in {".txt", ".md"}:
                    loaded = self._load_text_like_file(file_path)
                elif ext == ".docx":
                    loaded = self._load_docx(file_path)
                elif ext == ".csv":
                    loaded = self._load_csv(file_path)
                else:
                    loaded = []
            except Exception as e:
                print(f"Failed to load {file_path}: {e}")
                self.failed_files.append((file_path.name, f"{type(e).__name__}: {e}"))
                continue

            if not loaded:
                # File parsed without error but produced no extractable text — flag it
                # so callers can mark the DB row as failed rather than "indexed".
                self.failed_files.append((file_path.name, "no extractable text"))
                continue

            documents.extend(loaded)

        return documents

    def _load_pdf(self, pdf_file: Path) -> List[Document]:
        docs = []
        reader = PdfReader(pdf_file)
        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text()
            if text and text.strip():
                docs.append(
                    Document(
                        page_content=text,
                        metadata={"source": pdf_file.name, "page": page_num},
                    )
                )
        return docs

    def _load_text_like_file(self, file_path: Path) -> List[Document]:
        text = file_path.read_text(encoding="utf-8", errors="ignore")
        if not text.strip():
            return []
        return [
            Document(
                page_content=text,
                metadata={"source": file_path.name, "page": 1},
            )
        ]

    def _load_docx(self, file_path: Path) -> List[Document]:
        # Import lazily so the rest of the app still runs if python-docx is missing.
        from docx import Document as DocxDocument

        docx = DocxDocument(file_path)
        text_parts = [p.text.strip() for p in docx.paragraphs if p.text and p.text.strip()]
        full_text = "\n".join(text_parts)
        if not full_text:
            return []
        return [
            Document(
                page_content=full_text,
                metadata={"source": file_path.name, "page": 1},
            )
        ]

    def _load_csv(self, file_path: Path) -> List[Document]:
        docs = []
        with file_path.open("r", encoding="utf-8", errors="ignore", newline="") as f:
            reader = csv.reader(f)
            rows = list(reader)

        if not rows:
            return docs

        headers = rows[0]
        data_rows = rows[1:] if len(rows) > 1 else []

        for row_num, row in enumerate(data_rows, start=2):
            cells = []
            for idx, value in enumerate(row):
                header = headers[idx] if idx < len(headers) else f"column_{idx + 1}"
                cells.append(f"{header}: {value}")

            row_text = "\n".join(cells).strip()
            if row_text:
                docs.append(
                    Document(
                        page_content=row_text,
                        metadata={"source": file_path.name, "page": row_num},
                    )
                )

        # Fallback for header-only CSV
        if not docs:
            header_text = ", ".join(headers).strip()
            if header_text:
                docs.append(
                    Document(
                        page_content=f"CSV headers: {header_text}",
                        metadata={"source": file_path.name, "page": 1},
                    )
                )

        return docs

    def _build_kb_manifest(self) -> Dict[str, Any]:
        files = []
        supported_exts = {".pdf", ".txt", ".md", ".docx", ".csv"}

        for file_path in sorted(self.kb_path.iterdir(), key=lambda p: p.name.lower()):
            if not file_path.is_file() or file_path.suffix.lower() not in supported_exts:
                continue

            files.append(
                {
                    "name": file_path.name,
                    "size": file_path.stat().st_size,
                    "sha256": self._file_sha256(file_path),
                }
            )

        combined_hash = hashlib.sha256(
            json.dumps(files, sort_keys=True).encode("utf-8")
        ).hexdigest()
        return {"version": 1, "files": files, "combined_hash": combined_hash}

    def _file_sha256(self, file_path: Path) -> str:
        hasher = hashlib.sha256()
        with file_path.open("rb") as f:
            for chunk in iter(lambda: f.read(1024 * 1024), b""):
                hasher.update(chunk)
        return hasher.hexdigest()

    def _can_use_cached_index(self, current_manifest: Dict[str, Any]) -> bool:
        if not self.vector_store_path.exists() or not self.manifest_path.exists():
            return False

        try:
            saved_manifest = json.loads(self.manifest_path.read_text(encoding="utf-8"))
        except Exception:
            return False

        return saved_manifest.get("combined_hash") == current_manifest.get("combined_hash")

    def _write_manifest(self, manifest: Dict[str, Any]) -> None:
        self.manifest_path.write_text(
            json.dumps(manifest, indent=2),
            encoding="utf-8",
        )

    # -------------------------
    # Personal context
    # -------------------------
    def build_personal_context(
        self,
        user_id: int | None,
        user_message: str,
    ) -> Tuple[str, List[str]]:
        """Return (profile_text, memory_texts) for the given user.

        Safe to call with ``user_id=None`` (anonymous) or when profile/memory
        infrastructure isn't ready — returns empty values on any failure."""
        if not user_id:
            return "", []

        from .models import UserProfile
        from .memory import get_memory_store

        profile_text = ""
        try:
            profile = UserProfile.objects.get(user_id=user_id)
            parts: List[str] = []
            if profile.bio:
                parts.append(profile.bio.strip())
            prefs = profile.preferences or {}
            if prefs:
                pref_str = ", ".join(f"{k}: {v}" for k, v in prefs.items() if v)
                if pref_str:
                    parts.append(f"Preferences — {pref_str}")
            profile_text = " ".join(parts)
        except UserProfile.DoesNotExist:
            pass
        except Exception:
            pass

        memory_texts: List[str] = []
        try:
            store = get_memory_store()
            memory_texts = store.search(user_id, user_message, k=3)
        except Exception:
            pass

        return profile_text, memory_texts

    # -------------------------
    # Chat
    # -------------------------
    def chat(
        self,
        user_message: str,
        chat_history: List[Dict],
        user_id: int | None = None,
    ) -> Tuple[str, List[str], Dict[str, Any]]:
        """Returns (answer, sources, usage). `usage` is {tokens_in, tokens_out, cost_usd}
        — captures LLM call only; embedding cost is tracked separately if needed."""

        formatted_history = "\n".join(
            f"{m['type'].capitalize()}: {m['content']}"
            for m in chat_history[-4:]
        )

        intent = classify_intent(user_message)
        if intent != "factual":
            return self._chat_direct(user_message, formatted_history, user_id, intent)

        profile_text, memory_texts = self.build_personal_context(user_id, user_message)

        prompt = ChatPromptTemplate.from_template(CHAT_PROMPT_TEMPLATE)

        document_chain = create_stuff_documents_chain(
            self.llm,
            prompt
        )

        retrieval_chain = create_retrieval_chain(
            self.retriever,
            document_chain
        )

        try:
            with get_openai_callback() as cb:
                result = retrieval_chain.invoke(
                    {
                        "input": user_message,
                        "chat_history": formatted_history,
                        "user_profile": _format_profile(profile_text),
                        "user_memories": _format_memory_block(memory_texts),
                    }
                )

            answer = result["answer"]
            context_docs = result["context"]

            sources = [] if self._is_no_answer(answer) else self._extract_sources(context_docs)
            usage = {
                "tokens_in": cb.prompt_tokens,
                "tokens_out": cb.completion_tokens,
                "cost_usd": round(cb.total_cost, 6),
            }

            return answer, sources, usage

        except (APITimeoutError, APIConnectionError) as e:
            import traceback
            traceback.print_exc()
            raise UpstreamUnavailable(UPSTREAM_ERROR_MSG) from e

    def _chat_direct(
        self,
        user_message: str,
        formatted_history: str,
        user_id: int | None,
        intent: str,
    ) -> Tuple[str, List[str], Dict[str, Any]]:
        """LLM-only reply for conversational or personal-context turns. No retrieval."""
        template = PERSONAL_PROMPT_TEMPLATE if intent == "personal" else CONVERSATIONAL_PROMPT_TEMPLATE
        inputs: Dict[str, Any] = {"input": user_message, "chat_history": formatted_history}
        if intent == "personal":
            profile_text, memory_texts = self.build_personal_context(user_id, user_message)
            inputs["user_profile"] = _format_profile(profile_text)
            inputs["user_memories"] = _format_memory_block(memory_texts)

        chain = ChatPromptTemplate.from_template(template) | self.llm
        try:
            with get_openai_callback() as cb:
                resp = chain.invoke(inputs)
            answer = resp.content if hasattr(resp, "content") else str(resp)
            usage = {
                "tokens_in": cb.prompt_tokens,
                "tokens_out": cb.completion_tokens,
                "cost_usd": round(cb.total_cost, 6),
            }
            return answer, [], usage
        except (APITimeoutError, APIConnectionError) as e:
            import traceback
            traceback.print_exc()
            raise UpstreamUnavailable(UPSTREAM_ERROR_MSG) from e

    def extract_memories(
        self,
        user_message: str,
        assistant_reply: str,
    ) -> List[str]:
        """Extract durable third-person facts about the user from one exchange.

        Cheap gpt-4o-mini pass with a strict prompt. Returns ``[]`` on any error
        or when nothing worth remembering was shared — never raises. Results are
        post-filtered for obvious PII as a safety net in case the model ignores
        the instruction.
        """
        from .memory import contains_pii  # late import — avoid cycle at module load

        if not user_message or not assistant_reply:
            return []

        prompt = (
            "You extract DURABLE facts about the USER from one exchange, to personalize future chats.\n\n"
            "Return a JSON array of strings, nothing else. No prose, no markdown fences.\n\n"
            "EXTRACT only:\n"
            "- Identity / role (e.g. \"User is a data scientist at Acme\")\n"
            "- Projects, tools, companies the user explicitly mentions working with\n"
            "- Stable preferences (tone, format, expertise level)\n"
            "- Long-term goals or constraints the user states about themselves\n\n"
            "NEVER extract:\n"
            "- Transient state (what they asked just now, what they're currently doing)\n"
            "- Claims the ASSISTANT made — only what the user said about themselves\n"
            "- Passwords, API keys, credit card numbers, SSNs, medical diagnoses, precise addresses\n\n"
            "Format: 3rd person, one fact per string, <= 15 words each.\n"
            "If nothing durable was shared, return exactly: []\n\n"
            f"User: {user_message[:2000]}\n"
            f"Assistant: {assistant_reply[:2000]}\n\n"
            "JSON array:"
        )

        try:
            resp = self.llm.invoke(prompt)
            text = (resp.content if hasattr(resp, "content") else str(resp)).strip()
            # Strip markdown fences defensively if the model ignored instructions.
            if text.startswith("```"):
                text = text.strip("`").strip()
                if text.lower().startswith("json"):
                    text = text[4:].strip()
            facts = json.loads(text)
        except (APITimeoutError, APIConnectionError, json.JSONDecodeError, Exception):
            return []

        if not isinstance(facts, list):
            return []

        cleaned: List[str] = []
        for fact in facts:
            if not isinstance(fact, str):
                continue
            fact = fact.strip()
            if not fact or len(fact) > 240:
                continue
            if contains_pii(fact):
                continue
            cleaned.append(fact)
        return cleaned

    def generate_title(self, user_message: str, assistant_reply: str) -> str:
        """Produce a short (<=6 word) title from the first exchange.
        Returns empty string on any failure — caller falls back to message preview."""
        prompt = (
            "Summarize this conversation in a short descriptive title of 3-6 words. "
            "Return ONLY the title, no quotes, no punctuation at the end.\n\n"
            f"User: {user_message[:500]}\n"
            f"Assistant: {assistant_reply[:500]}"
        )
        try:
            resp = self.llm.invoke(prompt)
            text = (resp.content if hasattr(resp, 'content') else str(resp)).strip()
            # Strip quotes/trailing punctuation defensively
            text = text.strip('"\'').rstrip('.').strip()
            return text[:60]
        except Exception:
            return ""

    def chat_stream(
        self,
        user_message: str,
        chat_history: List[Dict],
        user_id: int | None = None,
    ):
        """
        Stream the RAG response token-by-token.
        Yields (content_chunk, None) for text, then (None, sources) at the end.
        """
        formatted_history = "\n".join(
            f"{m['type'].capitalize()}: {m['content']}"
            for m in chat_history[-4:]
        )

        intent = classify_intent(user_message)
        if intent != "factual":
            yield from self._chat_stream_direct(user_message, formatted_history, user_id, intent)
            return

        profile_text, memory_texts = self.build_personal_context(user_id, user_message)

        prompt = ChatPromptTemplate.from_template(CHAT_PROMPT_TEMPLATE)

        document_chain = create_stuff_documents_chain(
            self.llm_streaming,
            prompt,
        )

        try:
            # Get context first (no streaming)
            documents = self.retriever.invoke(user_message)
            sources = self._extract_sources(documents)

            inputs = {
                "input": user_message,
                "chat_history": formatted_history,
                "context": documents,
                "user_profile": _format_profile(profile_text),
                "user_memories": _format_memory_block(memory_texts),
            }

            answer_parts = []
            with get_openai_callback() as cb:
                for chunk in document_chain.stream(inputs):
                    if isinstance(chunk, dict) and "answer" in chunk:
                        part = chunk["answer"]
                    elif hasattr(chunk, "content"):
                        part = chunk.content
                    else:
                        part = str(chunk) if chunk else ""
                    if part:
                        answer_parts.append(part)
                        yield ("chunk", part)

            usage = {
                "tokens_in": cb.prompt_tokens,
                "tokens_out": cb.completion_tokens,
                "cost_usd": round(cb.total_cost, 6),
            }
            yield ("usage", usage)

            final_answer = "".join(answer_parts)
            final_sources = [] if self._is_no_answer(final_answer) else sources
            yield ("done", final_sources)

        except (APITimeoutError, APIConnectionError):
            import traceback
            traceback.print_exc()
            yield ("error", UPSTREAM_ERROR_MSG)
        except Exception:
            import traceback
            traceback.print_exc()
            yield ("error", GENERIC_ERROR_MSG)

    def _chat_stream_direct(
        self,
        user_message: str,
        formatted_history: str,
        user_id: int | None,
        intent: str,
    ):
        """Streaming LLM-only reply for conversational or personal-context turns. No retrieval."""
        template = PERSONAL_PROMPT_TEMPLATE if intent == "personal" else CONVERSATIONAL_PROMPT_TEMPLATE
        inputs: Dict[str, Any] = {"input": user_message, "chat_history": formatted_history}
        if intent == "personal":
            profile_text, memory_texts = self.build_personal_context(user_id, user_message)
            inputs["user_profile"] = _format_profile(profile_text)
            inputs["user_memories"] = _format_memory_block(memory_texts)

        chain = ChatPromptTemplate.from_template(template) | self.llm_streaming
        try:
            answer_parts = []
            with get_openai_callback() as cb:
                for chunk in chain.stream(inputs):
                    part = chunk.content if hasattr(chunk, "content") else (str(chunk) if chunk else "")
                    if part:
                        answer_parts.append(part)
                        yield ("chunk", part)
            usage = {
                "tokens_in": cb.prompt_tokens,
                "tokens_out": cb.completion_tokens,
                "cost_usd": round(cb.total_cost, 6),
            }
            yield ("usage", usage)
            yield ("done", [])
        except (APITimeoutError, APIConnectionError):
            import traceback
            traceback.print_exc()
            yield ("error", UPSTREAM_ERROR_MSG)
        except Exception:
            import traceback
            traceback.print_exc()
            yield ("error", GENERIC_ERROR_MSG)

    # -------------------------
    # Source citation
    # -------------------------
    @staticmethod
    def _is_no_answer(answer: str) -> bool:
        if not answer:
            return True
        normalized = answer.strip().rstrip(".").strip().lower()
        return normalized == NO_ANSWER_FALLBACK.strip().rstrip(".").strip().lower()

    def _extract_sources(
        self,
        docs: List[Document]
    ) -> List[str]:

        citations = []
        seen = set()

        for doc in docs:
            source = doc.metadata.get("source", "Unknown")
            page = doc.metadata.get("page", "?")
            citation = f"{source} - Page {page}"

            if citation not in seen:
                citations.append(citation)
                seen.add(citation)

        return citations


# -------------------------
# Singleton instance
# -------------------------
_rag_instance = None


def get_rag_system() -> RAGSystem:
    global _rag_instance
    if _rag_instance is None:
        _rag_instance = RAGSystem()
    return _rag_instance


def reload_rag_system() -> RAGSystem:
    """Force rebuild/load of the singleton RAG instance."""
    global _rag_instance
    _rag_instance = RAGSystem()
    return _rag_instance